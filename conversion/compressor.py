# conversion/compressor.py
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
from pathlib import Path
import os
from utils.file_handler import FileHandler


class DocumentCompressor:
    LEVELS = {"low": 0.5, "medium": 0.3, "high": 0.1}

    def compress_pdf(self, input_path: str, output_path: str, level: str = "medium") -> str:
        FileHandler.validate_file_exists(input_path)
        FileHandler.validate_file_extension(input_path, ['.pdf'])

        zoom = self.LEVELS.get(level, 0.3)
        mat = fitz.Matrix(zoom, zoom)
        
        # Buka PDF dengan mode bypass
        try:
            doc = fitz.open(input_path)
        except:
            print("[INFO] PDF terkunci! Membuka dengan mode bypass...")
            doc = fitz.open(filename=input_path, filetype="pdf", relaxed=True)

        # Unlock paksa kalau masih encrypted (99% PDF kuliah)
        if doc.is_encrypted:
            print("[INFO] PDF terenkripsi! Mencoba decrypt otomatis...")
            if doc.authenticate("") == 0:
                try:
                    doc._deleteObject(doc._xref_get_key(doc._xref_len() - 1)[1])
                    doc._updateStreamLengths()
                    print("[SUKSES] PDF berhasil di-unlock secara paksa!")
                except:
                    print("[GAGAL] Tidak bisa unlock. Coba print to PDF dulu.")

        total_pages = len(doc)
        print(f"Kompresi PDF: {total_pages} halaman, level={level}...")

        out_doc = fitz.open()
        
        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("jpeg", jpg_quality=45 if level == "high" else 65)
            new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
            new_page.insert_image(new_page.rect, stream=img_bytes)

        # SIMPAN TANPA linear=True → FIX ERROR CODE 4
        out_doc.save(
            output_path,
            garbage=4,
            deflate=True,
            clean=True,
            no_new_id=True
            # linear=True → SUDAH TIDAK SUPPORT LAGI!
        )
        out_doc.close()
        doc.close()

        size_in = os.path.getsize(input_path)
        size_out = os.path.getsize(output_path)
        reduction = 100 * (1 - size_out / size_in)
        print(f"[SELESAI] PDF terkompres → {output_path}")
        print(f"   Ukuran: {size_in/1024:.1f} KB → {size_out/1024:.1f} KB (-{reduction:.1f}%)")
        return output_path

    def compress_docx(self, input_path: str, output_path: str, level: str = "medium") -> str:
        FileHandler.validate_file_exists(input_path)
        FileHandler.validate_file_extension(input_path, ['.docx'])

        quality = 45 if level == "high" else 65 if level == "medium" else 85
        doc = Document(input_path)
        new_doc = Document()

        for para in doc.paragraphs:
            p = new_doc.add_paragraph(para.text)
            p.style = para.style

        for rel in doc.part.related_parts.values():
            if "image" not in rel.content_type:
                continue
            try:
                img_data = rel.blob
                img = Image.open(io.BytesIO(img_data))
                img_io = io.BytesIO()
                img = img.convert("RGB")
                img.save(img_io, format='JPEG', quality=quality, optimize=True, subsampling=2)
                img_io.seek(0)

                p = new_doc.add_paragraph()
                run = p.add_run()
                run.add_picture(img_io, width=Inches(5.5))
                p.alignment = 1
            except Exception as e:
                print(f"[WARN] Gagal kompres gambar: {e}")

        new_doc.save(output_path)
        print(f"[OK] DOCX terkompres → {output_path}")
        return output_path

    def compress(self, input_path: str, output_path: str, level: str = "medium") -> str:
        ext = Path(input_path).suffix.lower()
        if ext == ".pdf":
            return self.compress_pdf(input_path, output_path, level)
        elif ext == ".docx":
            return self.compress_docx(input_path, output_path, level)
        else:
            raise ValueError(f"Format tidak didukung: {ext}")