import os
import tempfile
import shutil
from abc import ABC, abstractmethod

try:
    from docx2pdf import convert as docx2pdf_convert
    LIBRARY_AVAILABLE = True
except ImportError:
    LIBRARY_AVAILABLE = False

try:
    import comtypes.client
    COMTYPES_AVAILABLE = True
except ImportError:
    COMTYPES_AVAILABLE = False

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdf2docx import Converter
    PDF2DOCX_CONVERTER_AVAILABLE = True
except ImportError:
    PDF2DOCX_CONVERTER_AVAILABLE = False


class ConversionStrategy(ABC):
    @abstractmethod
    def convert(self, input_file: str, output_file: str) -> bool: pass
    @abstractmethod
    def validate_input(self, input_file: str) -> bool: pass


class DocToPdfStrategy(ConversionStrategy):
    def __init__(self, has_ms_word: bool = False):
        self.has_ms_word = has_ms_word
        self._verified_ms_word = False

    def _verify_ms_word(self) -> bool:
        if not COMTYPES_AVAILABLE: return False
        try:
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            word.Quit()
            return True
        except: return False

    def validate_input(self, input_file: str) -> bool:
        if not os.path.exists(input_file):
            raise FileNotFoundError(f"File tidak ada: {input_file}")
        ext = os.path.splitext(input_file)[1].lower()
        if ext not in ['.doc', '.docx']:
            raise ValueError(f"Format salah: {ext}")
        if ext == '.doc' and not self.has_ms_word:
            raise ValueError("DOC butuh MS Word")
        if ext == '.doc' and not self._verified_ms_word:
            if not self._verify_ms_word():
                raise ValueError("MS Word tidak terdeteksi")
            self._verified_ms_word = True
        return True

    def convert(self, input_file: str, output_file: str) -> bool:
        self.validate_input(input_file)
        ext = os.path.splitext(input_file)[1].lower()
        if ext == '.docx':
            if not LIBRARY_AVAILABLE:
                raise ImportError("Install: pip install docx2pdf")
            docx2pdf_convert(input_file, output_file)
            return True
        else:
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(input_file))
            doc.SaveAs(os.path.abspath(output_file), FileFormat=17)
            doc.Close()
            word.Quit()
            return True


class PdfToDocxStrategy(ConversionStrategy):
    def __init__(self, method: str = "auto"):
        self.method = method

    def validate_input(self, input_file: str) -> bool:
        if not input_file.lower().endswith('.pdf'):
            raise ValueError("Harus PDF")
        if not os.path.exists(input_file):
            raise FileNotFoundError("File PDF tidak ada")
        return True

    def convert(self, input_file: str, output_file: str) -> bool:
        self.validate_input(input_file)

        if self.method in ["auto", "pdf2docx"] and PDF2DOCX_CONVERTER_AVAILABLE:
            try:
                cv = Converter(input_file)
                cv.convert(output_file)
                cv.close()
                return os.path.exists(output_file)
            except: pass

        if self.method in ["auto", "pymupdf"] and PYMUPDF_AVAILABLE and DOCX_AVAILABLE:
            try:
                return self._pymupdf(input_file, output_file)
            except: pass

        if self.method in ["auto", "text_only"] and PYMUPDF_AVAILABLE:
            try:
                return self._text_only(input_file, output_file)
            except: pass

        raise Exception("Semua metode gagal. Install pdf2docx")

    def _pymupdf(self, input_file: str, output_file: str) -> bool:
        temp_dir = tempfile.mkdtemp()
        try:
            pdf = fitz.open(input_file)
            doc = Document()
            for page in pdf:
                text = page.get_text()
                if text.strip():
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                for img in page.get_images():
                    xref = img[0]
                    pix = fitz.Pixmap(pdf, xref)
                    if pix.n - pix.alpha < 4:
                        img_path = os.path.join(temp_dir, f"img_{xref}.png")
                        pix.save(img_path)
                        doc.add_picture(img_path, width=Inches(6))
                        doc.add_paragraph()
                    pix = None
                doc.add_page_break()
            doc.save(output_file)
            return True
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def _text_only(self, input_file: str, output_file: str) -> bool:
        pdf = fitz.open(input_file)
        doc = Document()
        for page in pdf:
            text = page.get_text()
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.add_page_break()
        doc.save(output_file)
        pdf.close()
        return True


class PdfToDocStrategy(ConversionStrategy):
    def __init__(self, pdf_to_docx_strategy: PdfToDocxStrategy, has_ms_word: bool = False):
        self.pdf_to_docx = pdf_to_docx_strategy
        self.has_ms_word = has_ms_word

    def validate_input(self, input_file: str) -> bool:
        if not self.has_ms_word:
            raise Exception("PDF → DOC butuh MS Word")
        if not input_file.lower().endswith('.pdf'):
            raise ValueError("Input harus PDF")
        return True

    def convert(self, input_file: str, output_file: str) -> bool:
        self.validate_input(input_file)
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
        try:
            self.pdf_to_docx.convert(input_file, temp_docx)
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(temp_docx)
            doc.SaveAs(output_file, FileFormat=0)
            doc.Close()
            word.Quit()
            return True
        finally:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)